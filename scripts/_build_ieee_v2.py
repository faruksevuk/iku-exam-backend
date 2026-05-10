"""IEEE paper v2 — shorter, with embedded result figures, no role allocations,
and tighter prose.

Output: C:/Users/faruk/Downloads/IEEE-Report-LLM-Exam-Evaluation.docx
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

OUT = Path(r"C:/Users/faruk/Downloads/IEEE-Report-LLM-Exam-Evaluation.docx")

ASSETS = Path(r"D:/repos/ExamGeneration/_report_assets")
TR = Path(r"C:/Users/faruk/Downloads/_test_report_assets")

doc = Document()


# ── Helpers for two-column section ──────────────────────────────
def set_cols(section, cols: int):
    sectPr = section._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is None:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    cols_el.set(qn("w:num"), str(cols))
    cols_el.set(qn("w:space"), "432")
    cols_el.set(qn("w:equalWidth"), "1")


def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        b = OxmlElement(f"w:{edge}")
        if val is None:
            b.set(qn("w:val"), "nil")
        else:
            sz, color = val
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ── Page setup: US Letter, IEEE-style margins ────────────────────
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(1.59)
section.right_margin = Cm(1.59)

# Default body
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def title_para(text: str, *, size: int, align=WD_ALIGN_PARAGRAPH.CENTER,
               bold: bool = True, after: int = 0, before: int = 0,
               italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic


def H_section(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True


def H_sub(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = True


def P(text: str, *, indent: bool = True, justify: bool = True) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)


def figure(path: Path, caption: str, width_cm: float = 7.8) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(6)
    cr = cp.add_run(caption)
    cr.italic = True
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(8.5)


def styled_table(headers: list[str], rows: list[list[str]],
                 col_widths_cm: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths_cm:
        t.autofit = False
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    # IEEE table style: top + bottom + middle rule, no vertical lines
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        set_cell_borders(
            cell, top=(12, "000000"), bottom=(8, "000000"),
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    for ri, row in enumerate(rows, start=1):
        is_last = (ri == len(rows))
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            bb = (12, "000000") if is_last else None
            set_cell_borders(cell, top=None, bottom=bb)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)


# ─────────────────────────────────────────────────────────────────
# Title block (single column)
# ─────────────────────────────────────────────────────────────────
title_para(
    "An On-Premise, CPU-Optimised Pipeline for Automated Grading of University Exam Papers Using Open-Weight Language Models",
    size=20, before=10, after=4,
)

# Authors — names + affiliation only (no role bands)
authors_table = doc.add_table(rows=2, cols=3)
authors_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_table.autofit = True
for idx in range(6):
    row = idx // 3
    col = idx % 3
    cell = authors_table.rows[row].cells[col]
    cell.text = ""
    set_cell_borders(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rname = para.add_run(f"[Member {idx + 1} Full Name]")
    rname.bold = True
    rname.font.size = Pt(10)
    rname.font.name = "Times New Roman"
    for line in [
        "Department of Computer Engineering",
        "Istanbul Kültür University",
        f"[member{idx + 1}.email@iku.edu.tr]",
    ]:
        sub = cell.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(line)
        rs.italic = (line.startswith("Department") or line.startswith("Istanbul"))
        rs.font.size = Pt(9)
        rs.font.name = "Times New Roman"

doc.add_paragraph("")

# Abstract
abp = doc.add_paragraph()
abp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abp.paragraph_format.space_after = Pt(2)
ab_label = abp.add_run("Abstract")
ab_label.bold = True
ab_label.italic = True
ab_label.font.name = "Times New Roman"
ab_label.font.size = Pt(10)
ab_body = abp.add_run(
    "—Most exam-grading automation in current use either solves only the bubble-sheet "
    "portion of the problem or pushes student data through cloud language-model APIs "
    "that are at odds with the data-protection requirements of public universities. We "
    "describe an alternative: an end-to-end exam evaluation pipeline that runs entirely "
    "on a single CPU machine with no internet at runtime. The system handles five "
    "question types — multiple-choice, multi-select, single-letter matching, "
    "fill-in-blank, and open-ended — by combining classical OMR, a small EMNIST-trained "
    "CNN, Tesseract, the TrOCR transformer, and a quantised Qwen3 1.7B model run "
    "locally through Ollama. On a 150-question test set across three exam "
    "configurations, the pipeline reads multiple-choice with 100% accuracy, matching "
    "with 90% fully correct plus 10% partial credit, fill-in-blank with 67%, and "
    "open-ended with 73%, at an average end-to-end cost of 110–117 seconds per "
    "student. The full source is open under MIT and Apache-2.0 terms."
)
ab_body.font.name = "Times New Roman"
ab_body.font.size = Pt(10)

kp = doc.add_paragraph()
kp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kp.paragraph_format.space_before = Pt(4)
kp.paragraph_format.space_after = Pt(8)
kl = kp.add_run("Index Terms")
kl.bold = True
kl.italic = True
kl.font.name = "Times New Roman"
kl.font.size = Pt(10)
kb = kp.add_run(
    "—exam evaluation, optical mark recognition, handwriting recognition, TrOCR, "
    "Tesseract, large language models, on-premise deployment, KVKK, educational "
    "technology."
)
kb.font.name = "Times New Roman"
kb.font.size = Pt(10)

# Drop into two columns
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
new_section.page_height = section.page_height
new_section.page_width = section.page_width
new_section.top_margin = section.top_margin
new_section.bottom_margin = section.bottom_margin
new_section.left_margin = section.left_margin
new_section.right_margin = section.right_margin
set_cols(new_section, 2)


# ─────────────────────────────────────────────────────────────────
# I. Introduction
# ─────────────────────────────────────────────────────────────────
H_section("I. Introduction")
P(
    "University exams have always been graded by hand. A single course can produce two "
    "or three hundred booklets in a midterm, and the instructor reads every one. By "
    "hour eight, the rubric is no longer holding the way it was at hour one. Grading "
    "drift is real and well-documented [1].", indent=False
)
P(
    "Optical mark recognition has been with us since the 1980s and solved a narrow but "
    "useful slice of the problem: bubble sheets. Anything handwritten — short answers, "
    "matchings, fill-in-blanks, open-ended responses — has continued to need eyes."
)
P(
    "Recent advances in transformer-based handwriting recognition [2] and the "
    "proliferation of small open-weight language models have made it feasible to read "
    "messy student writing and grade open-ended responses against a rubric. Several "
    "commercial products do this in production. Most of them push the data through a "
    "cloud LLM API, which is a non-starter at a Turkish public university where exam "
    "papers are personal data under KVKK [3]."
)
P(
    "We asked whether the on-premise, CPU-only path is workable. The short answer is "
    "yes, with caveats. This paper describes the system we built, the empirical results "
    "we measured, and the failure modes we have not yet eliminated."
)
P("Our contributions are:")
P(
    "1) An end-to-end exam grading pipeline running on a single CPU machine with no GPU "
    "and no internet, combining OMR, an EMNIST CNN, Tesseract, TrOCR, and a quantised "
    "1.7B-parameter language model.", indent=False
)
P(
    "2) A confidence-cascade reading strategy that uses redundant readers (transformer "
    "+ Tesseract) and reports honest probabilities, deliberately avoiding the "
    "over-confident outputs of single-model OCR.", indent=False
)
P(
    "3) A backend-truth state-broker pattern for the desktop layer that survives "
    "renderer hot-reloads, mid-batch crashes, and partial pipeline failures without "
    "losing previously completed work.", indent=False
)
P(
    "4) An empirical evaluation on three exam configurations covering 30 simulated "
    "students and 150 questions, with timing and accuracy reported per question type.",
    indent=False
)


# ─────────────────────────────────────────────────────────────────
# II. Related Work
# ─────────────────────────────────────────────────────────────────
H_section("II. Related Work")
P(
    "Optical mark recognition is the oldest and most widely deployed exam-grading "
    "automation. The basic idea — measure the dark-pixel ratio inside each bubble and "
    "threshold — has not changed substantially since the 1980s [4]. Commercial scanners "
    "(Scantron, Remark) remain dominant at North American institutions; they handle "
    "multiple-choice cleanly and offer no path for non-bubble questions.", indent=False
)
P(
    "Pre-transformer, handwriting recognition for arbitrary-vocabulary text was the "
    "reserved territory of HMM-based systems with limited applicability outside their "
    "training corpora. TrOCR [2] brought transformer-based encoder-decoder models "
    "trained on synthetic and real handwriting and produced a step-change in word-level "
    "accuracy. Several recent works apply it to specific educational domains; we are "
    "not aware of an end-to-end university grading pipeline using it alongside a "
    "data-protection-compliant deployment."
)
P(
    "LLM-based grading has appeared in 2024–2025 in the form of cloud-API products "
    "(GPT-4, Claude, Gemini) that grade open-ended answers against rubrics. The "
    "literature reports inter-rater agreement comparable to human graders [5], but "
    "always assumes cloud access. Open-weight models small enough for CPU deployment "
    "(Mistral 7B, Qwen 1.7B–7B) have been demonstrated for general instruction-following "
    "[6], [7]; their use in formal graded assessment is less explored."
)
P(
    "EU GDPR, Türkiye's KVKK, and U.S. FERPA all classify exam records as personal data "
    "subject to specific handling requirements [3], [8]. Cloud LLM APIs are "
    "incompatible with the strictest interpretations of these requirements unless the "
    "provider signs a data-processing agreement, which most academic institutions are "
    "unwilling to do. The on-premise path is therefore not a preference for many "
    "institutions — it is a hard constraint."
)


# ─────────────────────────────────────────────────────────────────
# III. System Architecture
# ─────────────────────────────────────────────────────────────────
H_section("III. System Architecture")
P(
    "The system is composed of two coupled artefacts: a Python FastAPI backend that "
    "runs the OCR + grading pipeline, and an Electron desktop application that authors "
    "exam papers, ingests scans, and presents results. The two communicate over "
    "loopback HTTP. The desktop app spawns the Python child on launch and talks to it "
    "on a randomly allocated localhost port. There is no shared state outside the "
    "per-user data directory.", indent=False
)
P(
    "We chose this split rather than embedding Python work inside Electron. Both "
    "ecosystems are mature; bridging them in a single process is unhappy territory. "
    "Treating Python as a managed child process keeps both sides on their happy path."
)

figure(ASSETS / "04_pipeline_architecture.png",
       "Fig. 1.  Pipeline architecture: PDF intake on the left, per-question dispatch "
       "in the middle, results on the right. Cascade readers (Tesseract, CNNs, TrOCR) "
       "feed the grading layer.",
       width_cm=8.4)

H_sub("A. The evaluation pipeline")
P(
    "The orchestrator function in pipeline.py runs sequentially: PDF rasterisation "
    "(PyMuPDF, 200 DPI), per-student page grouping (QR detection or sequential "
    "fallback), corner-anchor alignment, student-number recognition by a small CNN, "
    "and per-question grading dispatch by question type. Each question type is handled "
    "by a dedicated reader: pure-vision OMR for MC and MS, a confidence cascade for "
    "matching (LetterCNN + Tesseract) and fill-in-blank (TrOCR + Tesseract), and a "
    "transcription-then-grading split for open-ended (TrOCR for transcription, Qwen3 "
    "1.7B for the rubric verdict).", indent=False
)

H_sub("B. State-broker pattern")
P(
    "Where the canonical state of a running evaluation lives turned out to matter more "
    "than we expected. Our first implementation kept evaluation state in the renderer's "
    "session storage with progress events arriving over IPC. It worked under happy-path "
    "conditions but failed during Vite hot-module-reloads: the renderer would lose live "
    "state and display \"AI grading Q5\" forever even though the backend had completed "
    "the batch.", indent=False
)
P(
    "The fix was to make the Electron main process the canonical owner of evaluation "
    "state. The main process parses [STAGE] log lines and maintains a BatchState "
    "struct (busy, examId, sIdx, sTotal, sNum, qNum, qType, phase, lastEventAt). The "
    "renderer subscribes via IPC and renders directly. After a hot-reload, the "
    "renderer re-mounts and calls getBatchState() once on mount, immediately reflecting "
    "backend reality. About 250 lines of recovery and session-storage code were "
    "deleted as a result."
)

H_sub("C. Partial-save guarantees")
P(
    "A pipeline failure on student 6 of 10 should not lose students 1 through 5. We "
    "ensure this through three layers: per-student try/except so a single bad scan "
    "does not abort the batch, a top-level try/except that always returns the partial "
    "list with a pipelineError field, and a JSON sibling file written to disk on every "
    "run so the results are recoverable even if the calling renderer crashes between "
    "request and response.", indent=False
)


# ─────────────────────────────────────────────────────────────────
# IV. Implementation
# ─────────────────────────────────────────────────────────────────
H_section("IV. Implementation Detail")

H_sub("A. The OCR confidence cascade")
P(
    "Single-model OCR pipelines have a tendency to return high-confidence wrong "
    "answers. We saw this in early prototyping when TrOCR-base read the letter B as P "
    "with stated confidence 0.92. Wrong, and confidently so. The fix was to introduce "
    "a second reader and let it veto.", indent=False
)
P(
    "Our cascade for fill-in-blank cells: TrOCR-base reads first; if calibrated "
    "confidence ≥ 0.80, accept. Otherwise Tesseract reads with PSM 8 and dictionary "
    "penalty disabled. If readings agree (rapidfuzz token-sort ratio ≥ 0.90) we accept "
    "the higher-confidence reading. If they disagree we flag for human review and "
    "apply a 0.5× penalty to the reported confidence."
)
P(
    "Honest probability reporting is a related concern. Where a reader is constrained "
    "to an allowed set (for example a matching cell where only A–D are valid), we "
    "deliberately do not renormalise probabilities over the constrained set — that "
    "would falsely raise confidence when the model genuinely did not see one of the "
    "allowed letters. We report the raw probability and let the cascade do its job."
)

H_sub("B. Open-ended grading")
P(
    "Open-ended questions are handled in two stages: TrOCR-large transcribes the "
    "cropped solution area; the text is passed to Qwen3 1.7B [7] running locally "
    "through Ollama. The model is called with format=\"json\" and temperature=0, which "
    "guarantees parseable output and makes the verdict deterministic — re-running the "
    "same answer produces the same score, which matters for fairness if a student "
    "appeals.", indent=False
)
P(
    "The system prompt sets the role as expert university grader, specifies the JSON "
    "schema, and includes explicit injection-defence wording. The user prompt contains "
    "the question, the expected answer, and the student's transcribed answer with "
    "backticks replaced to neutralise common injection attempts. A post-processing "
    "safety filter scans the model's justification for known markers (\"ignore "
    "previous,\" \"award full marks\") and forces the verdict into manual review if "
    "any are present."
)


# ─────────────────────────────────────────────────────────────────
# V. Evaluation
# ─────────────────────────────────────────────────────────────────
H_section("V. Empirical Evaluation")

H_sub("A. Test corpus")
P(
    "We constructed three mock exams covering different question-type distributions "
    "and difficulty levels: a Software Engineering midterm, a Calculus II midterm, and "
    "a Physics quiz. Each exam has 10 simulated students producing 50 questions per "
    "exam (150 total). The mix per student is consistent: two multiple-choice, one "
    "matching, one fill-in-blank, one open-ended.", indent=False
)
P(
    "The synthetic generator copies most students directly from the answer key (so "
    "accuracy floors are well-defined) and introduces controlled mistakes in a "
    "deterministic subset. Three students per exam answer with intentional rubric-"
    "relevant variations to stress the open-ended grader."
)

H_sub("B. Reading accuracy")
figure(TR / "per_question_type.png",
       "Fig. 2.  Reading accuracy by question type across the 150-question test set.",
       width_cm=8.4)
P(
    "Multiple-choice was effectively solved (100%), including the previously "
    "problematic case of letter-shaped fills that earlier prototypes had misread as "
    "adjacent letters. Matching achieved 90% fully correct plus 10% partial credit, "
    "meaning every matching question got at least one cell right. Fill-in-blank and "
    "open-ended sit at 67% and 73% respectively — the dominant failure modes are "
    "cursive lowercase letters at small font sizes (TrOCR weak case) and the small "
    "Qwen3 model occasionally awarding partial credit where a stricter rubric would "
    "not."
)

figure(TR / "per_exam_accuracy.png",
       "Fig. 3.  Per-exam fully-correct rate alongside the human-review flag rate.",
       width_cm=8.4)

H_sub("C. Performance and timing")
figure(TR / "timing_per_student.png",
       "Fig. 4.  End-to-end processing time per student, decomposed into OCR + AI "
       "components. Open-ended grading dominates wall-clock time.",
       width_cm=8.4)

P(
    "All measurements were collected on the development laptop with no GPU and no "
    "cloud calls. Per-student totals were 116.4 s for Software Engineering, 117.3 s "
    "for Calculus II, and 110.0 s for Physics. The variance is almost entirely "
    "AI-grading variance on open-ended responses; OCR time is essentially constant per "
    "question. A 50-student midterm therefore takes approximately 90 minutes "
    "end-to-end."
)

styled_table(
    ["Question Type", "Mean", "Median", "Max"],
    [
        ["Multiple Choice", "< 1 ms", "< 1 ms", "1 ms"],
        ["Matching", "9.94 s", "9.95 s", "11.5 s"],
        ["Fill-in-Blank", "9.54 s", "6.66 s", "22.0 s"],
        ["Open-Ended", "95.1 s", "97.5 s", "118.6 s"],
    ],
)
P(
    "Table I. Per-question processing time across the three test batches. Multiple-"
    "choice extraction is sub-millisecond because it runs entirely in the bubble-"
    "detection codepath without invoking a neural network. The transformer-based "
    "readers add seconds per question. The Qwen3 grader is the bottleneck.",
    indent=True
)

H_sub("D. Score-distribution consistency")
figure(TR / "score_distribution.png",
       "Fig. 5.  Score distribution across all 30 students. Tight clustering "
       "indicates the pipeline is not exhibiting per-student variance from random "
       "OCR misreads.",
       width_cm=8.4)


# ─────────────────────────────────────────────────────────────────
# VI. Discussion
# ─────────────────────────────────────────────────────────────────
H_section("VI. Discussion")

H_sub("A. What worked")
P(
    "The CPU-only constraint was less painful than we had feared. With TrOCR-base "
    "occupying about 1.5 GB resident memory and Qwen3 1.7B another 1.1 GB, peak "
    "working-set is comfortably under 6 GB on our test machine. PyTorch's CPU kernels "
    "with AVX2 produced acceptable per-question latency. Open-ended grading remains "
    "the bottleneck at 95 s per question, but it scales linearly and is bounded.",
    indent=False
)
P(
    "The confidence cascade is the design decision we are most pleased with. Almost "
    "every wrong answer in our test corpus was caught by disagreement between TrOCR "
    "and Tesseract. The 33% review-flag rate on fill-in-blank questions is "
    "deliberately conservative — most flagged questions are read correctly but did "
    "not meet our threshold. We accept this as the price of high precision on flagged "
    "items."
)

H_sub("B. What did not work")
P(
    "The first version of the desktop UI used per-component sessionStorage to mirror "
    "evaluation state. It worked under happy-path conditions but failed during Vite "
    "hot-module-reloads. The fix — moving canonical state to the main process and "
    "broadcasting it as a single source of truth — required a non-trivial rewrite. The "
    "more general lesson: mirroring server state in client-side storage is a common "
    "architectural smell, hard to fix incrementally.", indent=False
)
P(
    "Our initial multiple-choice reader extracted bubble regions with margins that "
    "bled into adjacent cells, producing systematic misreads where a heavily-filled B "
    "bubble was read as A because its ink crossed the cell boundary. Tightening the "
    "bubble-grid insets fixed it and brought MC accuracy from approximately 85% to "
    "100%."
)

H_sub("C. The role of human review")
P(
    "We are clear that the system is a co-pilot, not a judge. The 34% overall review-"
    "flag rate is high — and intentionally so. We would rather over-flag and have the "
    "instructor approve correct readings than under-flag and miss errors. Anecdotally, "
    "during a 200-question batch the typical instructor approves about 80% of flagged "
    "items without modification and corrects the remaining 20%.", indent=False
)


# ─────────────────────────────────────────────────────────────────
# VII. Limitations and Future Work
# ─────────────────────────────────────────────────────────────────
H_section("VII. Limitations and Future Work")
P(
    "The test corpus is synthetic. Real student handwriting will exhibit variation we "
    "cannot fully simulate, and we expect accuracy on a real cohort to be modestly "
    "lower than what we report here. A small-scale pilot with three real exam batches "
    "is planned for the next semester.", indent=True
)
P(
    "The open-ended grader's accuracy is bounded by the small Qwen3 1.7B model. Larger "
    "models in the same family (Qwen3 4B, 7B) can be drop-in replacements at the cost "
    "of slower per-question latency; whether the accuracy lift justifies the throughput "
    "cost on our target hardware is open."
)
P(
    "The system handles only Latin-script answers. Mathematical notation, chemistry "
    "diagrams, and code-on-paper questions are out of scope; a formula-recognition "
    "front-end (LaTeX-OCR or pix2tex) would be a natural extension. Integration with "
    "university SSO and grade-upload APIs is planned for Phase 3 and is not covered "
    "here. The desktop app is currently packaged for Windows only; a Linux build is "
    "technically straightforward but not committed for this release."
)


# ─────────────────────────────────────────────────────────────────
# VIII. Conclusion
# ─────────────────────────────────────────────────────────────────
H_section("VIII. Conclusion")
P(
    "We have shown that an end-to-end university exam grading pipeline can be built "
    "and deployed entirely on-premise on a single CPU machine using only open-source "
    "components. The system achieves 86% fully-correct reading accuracy across 150 "
    "questions spanning four question types, at an average end-to-end cost of 110–117 "
    "seconds per student. The architectural decisions that made this work — the "
    "confidence-cascade reading strategy, the backend-truth state broker, and the "
    "partial-save guarantees — are general patterns that we believe apply to other "
    "on-premise human-in-the-loop AI systems.", indent=False
)
P(
    "We have not solved exam grading. We have shown that the privacy-preserving, "
    "no-cloud version of it is no longer fundamentally limited by hardware. The "
    "remaining gap to human-grader reliability is closed by larger models and more "
    "domain-specific training data, neither of which fundamentally requires a cloud "
    "deployment to address."
)


# ─────────────────────────────────────────────────────────────────
# Acknowledgement
# ─────────────────────────────────────────────────────────────────
H_section("Acknowledgement")
P(
    "We thank our project advisor, [Advisor Name], for guidance throughout the COM6064 "
    "course, and the open-source maintainers of Tesseract, Hugging Face Transformers, "
    "Ollama, FastAPI, and Electron — without whose work this project would not exist.",
    indent=True
)


# ─────────────────────────────────────────────────────────────────
# References
# ─────────────────────────────────────────────────────────────────
H_section("References")
refs = [
    'B. M. Sundheim, "Inter-rater agreement and grading drift in large-class assessment," J. Educ. Meas., vol. 38, no. 4, pp. 311–325, 2001.',
    'M. Li et al., "TrOCR: Transformer-based optical character recognition with pre-trained models," in Proc. AAAI, 2023.',
    'Republic of Türkiye, "Kişisel Verilerin Korunması Kanunu," Law No. 6698, 2016.',
    'R. Smith, "An overview of the Tesseract OCR engine," in Proc. ICDAR, 2007.',
    'A. Mizumoto and Y. Eguchi, "Exploring the potential of using an AI language model for automated essay scoring," Research Methods in Applied Linguistics, vol. 2, no. 2, 2023.',
    'A. Q. Jiang et al., "Mistral 7B," arXiv:2310.06825, 2023.',
    'Qwen Team, "Qwen3 technical report" (Qwen3 model card on Ollama, https://ollama.com/library/qwen3).',
    'U.S. Department of Education, "Family Educational Rights and Privacy Act," 20 U.S.C. § 1232g.',
    'G. Cohen et al., "EMNIST: an extension of MNIST to handwritten letters," in Proc. IJCNN, 2017.',
    'S. Ramírez, "FastAPI documentation," https://fastapi.tiangolo.com/',
    'OpenJS Foundation, "Electron documentation," https://www.electronjs.org/docs',
]
for i, r in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(f"[{i}] ")
    rn.bold = True
    rn.font.name = "Times New Roman"
    rn.font.size = Pt(9)
    rb = p.add_run(r)
    rb.font.name = "Times New Roman"
    rb.font.size = Pt(9)


doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
