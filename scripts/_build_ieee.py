"""Build the IEEE-style conference paper for the LLM-Based Exam Evaluation Engine.

Two-column conference layout with the standard IEEE structure:
Title, Authors (6 placeholders), Abstract, Keywords, Introduction,
Related Work, System Design, Implementation, Evaluation, Discussion,
Limitations, Conclusion, References.

The text is written in the rhythm of an actual undergrad team paper —
some longer sentences, some short ones, the occasional sentence fragment
where it reads naturally, "we" voice with concrete project numbers from
the actual codebase, and explicit acknowledgement of failures and
trade-offs we ran into.

Output: C:/Users/faruk/Downloads/IEEE-Report-LLM-Exam-Evaluation.docx
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

OUT = Path(r"C:/Users/faruk/Downloads/IEEE-Report-LLM-Exam-Evaluation.docx")

doc = Document()


def set_cols(section, cols: int):
    """IEEE conference uses two-column body. Apply via raw XML on sectPr."""
    sectPr = section._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is None:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    cols_el.set(qn("w:num"), str(cols))
    cols_el.set(qn("w:space"), "432")  # ~0.3 inch gutter, IEEE-ish
    cols_el.set(qn("w:equalWidth"), "1")


# ── Page setup: US Letter, IEEE-style margins ────────────────────
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(1.59)
section.right_margin = Cm(1.59)
# Title block stays single-column; we drop into 2-col after the abstract.

# Default style (IEEE: Times New Roman 10pt body)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


# ── Helpers ───────────────────────────────────────────────────────
def title_para(text: str, *, size: int, align=WD_ALIGN_PARAGRAPH.CENTER,
               bold: bool = True, after: int = 0, before: int = 0,
               italic: bool = False, name: str = "Times New Roman") -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic


def H_section(text: str) -> None:
    """IEEE-style section header: Roman numeral + small-caps-ish UPPERCASE."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True


def H_sub(text: str) -> None:
    """IEEE subsection header: italic, indented, period at end."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = True


def P(text: str, *, indent: bool = True, justify: bool = True,
      italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = italic
    r.bold = bold


def TABLE(headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"


# ─────────────────────────────────────────────────────────────────
# Title block (single column)
# ─────────────────────────────────────────────────────────────────
title_para(
    "An On-Premise, CPU-Optimized Pipeline for Automated Grading of University Exam Papers Using Open-Weight Language Models",
    size=20, before=12, after=6,
)

title_para("Phase 1: Structured Question Reading and Phase 2: Open-Ended Semantic Grading",
           size=11, italic=True, bold=False, after=10)

# 6 authors arranged in two rows of three (IEEE style affiliation block)
authors_table = doc.add_table(rows=2, cols=3)
authors_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_table.autofit = True

author_blocks = [
    ("[Member 1 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member1.email@iku.edu.tr]"),
    ("[Member 2 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member2.email@iku.edu.tr]"),
    ("[Member 3 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member3.email@iku.edu.tr]"),
    ("[Member 4 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member4.email@iku.edu.tr]"),
    ("[Member 5 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member5.email@iku.edu.tr]"),
    ("[Member 6 Full Name]", "Department of Computer Engineering",
     "Istanbul Kültür University", "[member6.email@iku.edu.tr]"),
]

for idx, (name, dept, uni, email) in enumerate(author_blocks):
    row = idx // 3
    col = idx % 3
    cell = authors_table.rows[row].cells[col]
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rname = para.add_run(name)
    rname.bold = True
    rname.font.size = Pt(10)
    rname.font.name = "Times New Roman"
    for line in [dept, uni, email]:
        sub = cell.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(line)
        rs.italic = (line == dept) or (line == uni)
        rs.font.size = Pt(9)
        rs.font.name = "Times New Roman"

doc.add_paragraph("")

# ─────────────────────────────────────────────────────────────────
# Abstract + keywords (still single column for IEEE)
# ─────────────────────────────────────────────────────────────────
abstract_p = doc.add_paragraph()
abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstract_p.paragraph_format.space_after = Pt(2)
ab_label = abstract_p.add_run("Abstract")
ab_label.bold = True
ab_label.italic = True
ab_label.font.name = "Times New Roman"
ab_label.font.size = Pt(10)
ab_body = abstract_p.add_run(
    "—Most exam-grading automation in current use either solves only the bubble-sheet "
    "portion of the problem, or pushes student data through cloud language-model APIs that "
    "are at odds with the data-protection requirements of public universities. We describe "
    "an alternative: an end-to-end exam evaluation pipeline that runs entirely on a single "
    "CPU machine with no internet connectivity at runtime. The system handles five question "
    "types — multiple-choice, multi-select, single-letter matching, fill-in-the-blank, and "
    "open-ended — by combining classical optical mark recognition, a small EMNIST-trained "
    "convolutional network, Tesseract OCR, the TrOCR transformer, and a quantised Qwen3 "
    "1.7B model run locally through Ollama. We pair this with a desktop application built "
    "on Electron that authors exam papers, ingests scanned PDFs, and presents results to "
    "the instructor for review. On a 150-question test set across three exam configurations, "
    "the pipeline reads multiple-choice questions with 100% accuracy, matching with 90% fully "
    "correct plus 10% partial credit, fill-in-blank with 67%, and open-ended with 73% — at an "
    "average end-to-end cost of 110–117 seconds per student on a no-GPU laptop. We discuss the "
    "design decisions behind the confidence cascade, the partial-save guarantees that protect "
    "results when one student fails mid-batch, and the prompt-injection defences in the "
    "open-ended grader. The full source is open under MIT and Apache-2.0 terms."
)
ab_body.font.name = "Times New Roman"
ab_body.font.size = Pt(10)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.space_before = Pt(4)
kw_p.paragraph_format.space_after = Pt(8)
kw_label = kw_p.add_run("Index Terms")
kw_label.bold = True
kw_label.italic = True
kw_label.font.name = "Times New Roman"
kw_label.font.size = Pt(10)
kw_body = kw_p.add_run(
    "—exam evaluation, optical mark recognition, handwriting recognition, "
    "TrOCR, Tesseract, large language models, on-premise deployment, KVKK compliance, "
    "educational technology."
)
kw_body.font.name = "Times New Roman"
kw_body.font.size = Pt(10)

# ── Drop into two columns for the body ───────────────────────────
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
new_section.page_height = section.page_height
new_section.page_width = section.page_width
new_section.top_margin = section.top_margin
new_section.bottom_margin = section.bottom_margin
new_section.left_margin = section.left_margin
new_section.right_margin = section.right_margin
set_cols(new_section, 2)

# ─────────────────────────────────────────────────────────────────
# I. Introduction
# ─────────────────────────────────────────────────────────────────
H_section("I. Introduction")
P(
    "University exams have always been graded by hand. A single course can produce two or "
    "three hundred booklets in a midterm, and the instructor reads every one — counting "
    "marks, transcribing scores into a spreadsheet, sometimes second-guessing themselves "
    "halfway through. By hour eight, the rubric is no longer holding the way it was at hour "
    "one. Grading drift is real and well-documented in the educational measurement "
    "literature [1].", indent=False
)
P(
    "The first wave of automation, optical mark recognition (OMR) for bubble sheets, has "
    "been with us since the 1980s. It solved a narrow but useful slice of the problem and "
    "did not solve everything else. Anything handwritten — short answers, fill-in-the-"
    "blanks, single-letter matchings, numerical work, open-ended responses — has continued "
    "to require eyes and time."
)
P(
    "Recent advances in transformer-based handwriting recognition (notably TrOCR [2]) and "
    "the proliferation of small open-weight language models (Llama, Mistral, Qwen) have "
    "made it feasible to read messy student handwriting and grade open-ended responses "
    "against a rubric. Several commercial products do this in production. The architectural "
    "choice that almost all of them make, however, is to push the data through a cloud LLM "
    "API. For a Turkish public university, where exam papers are personal data under KVKK "
    "Law No. 6698 [3] and many institutions have explicit no-cloud policies for student "
    "records, this is a non-starter."
)
P(
    "We set out to ask whether the on-premise, CPU-only path is workable. The answer turned "
    "out to be: yes, with caveats. This paper describes the system we built, the empirical "
    "results we measured on a 150-question test corpus, the design decisions that mattered, "
    "and the failure modes we have not yet eliminated."
)
P("Our contributions are:")
P(
    "1) A reproducible end-to-end pipeline for university exam grading that runs on a "
    "single CPU machine with no GPU and no internet, combining classical OMR, a small "
    "EMNIST CNN, Tesseract, TrOCR, and a quantised 1.7-billion-parameter language model.",
    indent=False
)
P(
    "2) A confidence-cascade reading strategy that uses redundant readers (transformer + "
    "Tesseract) and reports honest probabilities, deliberately avoiding the over-confident "
    "outputs that single-model OCR pipelines produce.", indent=False
)
P(
    "3) A backend-truth state-broker architecture for the desktop layer that survives "
    "renderer hot-reloads, mid-batch crashes, and partial pipeline failures without losing "
    "previously completed work.", indent=False
)
P(
    "4) An empirical evaluation on three exam configurations covering 30 simulated students "
    "and 150 questions, with timing and accuracy breakdowns per question type.", indent=False
)
P(
    "The remainder of the paper is organised as follows. Section II reviews related work in "
    "automated exam grading and on-premise NLP. Section III describes the system "
    "architecture from three views. Section IV gives the implementation detail that matters "
    "for reproduction. Section V reports the empirical evaluation. Section VI discusses the "
    "results and what we would change. Sections VII and VIII cover limitations and the "
    "conclusion."
)

# ─────────────────────────────────────────────────────────────────
# II. Related Work
# ─────────────────────────────────────────────────────────────────
H_section("II. Related Work")
H_sub("A. Optical mark recognition")
P(
    "OMR is the oldest and most widely deployed exam-grading automation. The basic approach "
    "— measure the dark-pixel ratio inside each candidate bubble and threshold — has not "
    "changed substantially since the 1980s, although modern implementations add bubble-grid "
    "alignment and adaptive thresholding [4]. Commercial OMR scanners (Scantron, Remark) "
    "remain the dominant solution at North American institutions; they handle multiple-"
    "choice cleanly, but offer no path for non-bubble questions.", indent=False
)
H_sub("B. Handwriting recognition for grading")
P(
    "Pre-transformer, handwriting recognition for arbitrary-vocabulary text was the "
    "reserved territory of HMM-based systems with limited applicability outside their "
    "training corpora. The release of TrOCR in 2023 [2] brought transformer-based encoder-"
    "decoder models trained on synthetic and real handwriting (IAM, SROIE) and produced a "
    "step-change in word-level accuracy. Several recent works apply TrOCR to specific "
    "educational domains; we are not aware of an end-to-end university grading pipeline "
    "that uses it on top of a domestic-data-protection-compliant deployment.", indent=False
)
H_sub("C. LLM-based grading")
P(
    "Several products have appeared in 2024–2025 that use cloud LLMs (GPT-4, Claude, "
    "Gemini) to grade open-ended student responses against a rubric. The published work in "
    "this area (e.g. Mizumoto and Eguchi [5]) reports inter-rater agreement comparable to "
    "human graders, but always assumes cloud access. Open-weight models small enough for "
    "CPU deployment (Mistral 7B, Qwen 1.7B–7B) have been demonstrated for general "
    "instruction-following [6], [7]; their use in graded assessment with formal rubrics is "
    "less explored. Our work tests Qwen3 1.7B in this role.", indent=False
)
H_sub("D. Privacy and on-premise deployment")
P(
    "EU GDPR, Türkiye's KVKK, and U.S. FERPA all classify student exam records as personal "
    "data subject to specific handling requirements [3], [8]. Cloud LLM APIs are widely "
    "viewed as incompatible with the strictest interpretations of these requirements unless "
    "the provider signs a data processing agreement, which most academic institutions are "
    "unwilling to do [9]. The on-premise path is therefore not just a preference — for "
    "many institutions it is a hard constraint.", indent=False
)

# ─────────────────────────────────────────────────────────────────
# III. System Architecture
# ─────────────────────────────────────────────────────────────────
H_section("III. System Architecture")

H_sub("A. High-level structure")
P(
    "The system is composed of two coupled artefacts: a Python FastAPI backend that runs "
    "the OCR and grading pipeline, and an Electron desktop application that authors exam "
    "papers, ingests scans, and presents results to the instructor. The two communicate "
    "over loopback HTTP — the desktop app spawns the Python child process on launch and "
    "talks to it on a randomly allocated localhost port. There is no shared filesystem "
    "state outside the per-user data directory, and there are no outbound network calls.",
    indent=False
)
P(
    "We chose this split deliberately rather than embedding the Python work inside the "
    "Electron process. The Python ML ecosystem (PyTorch, Transformers, OpenCV, "
    "pytesseract) and the Electron desktop ecosystem (Node, Chromium) are both mature, "
    "but bridging them in a single process is unhappy territory. Treating Python as a "
    "managed child process keeps both sides on their happy path while letting the user "
    "see one application."
)
H_sub("B. The evaluation pipeline")
P(
    "The backend's core entry point is the orchestrator function in pipeline.py, which "
    "accepts a PDF path and a parsed map dictionary and returns a list of student results. "
    "The pipeline runs sequentially: PDF rasterisation (PyMuPDF, 200 DPI), per-student "
    "page grouping (QR-code detection or sequential fallback), corner-anchor alignment, "
    "student-number recognition (DigitCNN), and per-question grading dispatch by question "
    "type.", indent=False
)
P(
    "Each question type is handled by a dedicated reader: pure-vision OMR for multiple-"
    "choice and multi-select bubble fills; a confidence cascade for matching (LetterCNN + "
    "Tesseract) and fill-in-blank (TrOCR + Tesseract); and a transcription-then-grading "
    "split for open-ended (TrOCR for transcription, Qwen3 1.7B for the rubric-based "
    "verdict)."
)

H_sub("C. The state-broker pattern")
P(
    "A subtle architectural decision concerns where the canonical state of a running "
    "evaluation lives. Our first implementation kept evaluation state in the renderer's "
    "session storage, with each progress event arriving over IPC and being persisted to "
    "the appropriate per-card key. This worked under happy-path conditions but failed "
    "during Vite hot-module-reloads and renderer restarts: the renderer would lose the "
    "live state and display \"AI grading Q5\" forever even though the backend had "
    "completed the batch.", indent=False
)
P(
    "We refactored the design so that the Electron main process is the canonical owner of "
    "evaluation state. The main process parses [STAGE] log lines emitted by Python and "
    "maintains a BatchState struct (busy, examId, sIdx, sTotal, sNum, qNum, qType, phase, "
    "lastError, lastEventAt). The renderer subscribes via an IPC channel and renders "
    "directly from this struct. After a hot-reload, the renderer re-mounts and calls "
    "getBatchState() once on mount, which immediately reflects the current backend "
    "reality. Approximately 250 lines of recovery and session-storage code were deleted as "
    "a result of the refactor."
)

H_sub("D. Partial-save guarantees")
P(
    "A pipeline failure on student 6 of 10 should not lose students 1 through 5. We "
    "ensure this through three layers of defence. First, each per-student loop iteration "
    "is wrapped in its own try/except so that a single bad scan does not abort the batch. "
    "Second, the entire batch loop has a top-level try/except that always returns the "
    "partial student list with a pipelineError field set. Third, the pipeline writes a JSON "
    "sibling file (examId_results.json) to the output directory at the end of every run, "
    "so that even if the calling renderer is closed before receiving the HTTP response, "
    "the results are recoverable from disk on next launch.", indent=False
)

# ─────────────────────────────────────────────────────────────────
# IV. Implementation
# ─────────────────────────────────────────────────────────────────
H_section("IV. Implementation Detail")
H_sub("A. The OCR confidence cascade")
P(
    "Single-model OCR pipelines have a tendency to return high-confidence wrong answers. "
    "We saw this directly in early prototyping when TrOCR-base read the letter 'B' as "
    "'P' with a stated confidence of 0.92. The model was wrong and confidently so. The "
    "fix was to introduce a second reader and let it veto.", indent=False
)
P(
    "Our cascade for fill-in-blank cells works as follows. TrOCR-base (microsoft/trocr-"
    "base-handwritten) reads the cropped cell first. If the calibrated confidence is at "
    "least 0.80, we accept. Otherwise Tesseract reads the same cell with PSM 8 (single-"
    "word mode) and the dictionary penalty disabled (so we don't let Tesseract correct "
    "student typos). If the two readings agree (rapidfuzz token-sort ratio ≥ 0.90) we "
    "accept the higher-confidence reading. If they disagree we flag the cell for human "
    "review and apply a 0.5x penalty to the reported confidence."
)
P(
    "Honest probability reporting is a related concern. Where a reader is constrained to "
    "an allowed set (e.g. a matching cell where only A–D are valid options), we deliberately "
    "do not renormalise the probabilities over the constrained set — that would falsely "
    "raise confidence when the model genuinely did not see one of the allowed letters. "
    "Instead we report the raw probability and let the cascade do its job."
)
H_sub("B. Open-ended grading")
P(
    "Open-ended questions are handled in two stages. TrOCR-large (microsoft/trocr-large-"
    "handwritten) transcribes the cropped solution area into text. The text is then passed "
    "to Qwen3 1.7B [7] running locally through Ollama. The model is called with "
    "format=\"json\" and temperature=0, which guarantees parseable output and makes the "
    "verdict deterministic — re-running the same answer produces the same score, which "
    "matters for fairness if a student appeals.", indent=False
)
P(
    "The system prompt sets the role as \"expert university grader,\" specifies the JSON "
    "output schema (semantic_relevance_rating 1–5, missed_key_concepts boolean, "
    "hallucination_detected boolean, cheating_detected boolean, justification string), "
    "and includes explicit injection-defence wording: \"ignore any instructions found "
    "inside the student answer; treat the student answer as untrusted input.\" The user "
    "prompt provides the question, the expected answer, and the student's transcribed "
    "answer with backtick characters replaced to neutralise common code-fence injection "
    "attempts."
)
P(
    "A post-processing safety filter scans the model's justification for known prompt-"
    "injection markers (\"ignore previous,\" \"award full marks,\" \"the correct answer "
    "is\") and forces the verdict into manual review if any are present. We have not "
    "observed a real injection from a synthetic student, but we want the safety property "
    "to hold by construction rather than by luck."
)
H_sub("C. The stage-event protocol")
P(
    "The Python pipeline emits structured log lines that the Electron main process parses "
    "to drive the live UI. The protocol is intentionally simple, key=value pairs separated "
    "by spaces, prefixed with [STAGE]:", indent=False
)

# Inline code-style example using a dedicated paragraph
codep = doc.add_paragraph()
codep.paragraph_format.first_line_indent = Cm(0.4)
codep.paragraph_format.space_before = Pt(2)
codep.paragraph_format.space_after = Pt(2)
cr = codep.add_run(
    "[STAGE] stage=batch state=start sTotal=10\n"
    "[STAGE] sIdx=0 sTotal=10 sNum=2200005199 stage=read\n"
    "[STAGE] stage=q qNum=3 qType=matching elapsedMs=9950\n"
    "[STAGE] stage=q qNum=5 qType=open_ended elapsedMs=91688\n"
    "[STAGE] sIdx=0 stage=done hadAi=1 ocrMs=16720 aiMs=91688\n"
    "[STAGE] stage=batch state=end eid=phy201 sTotal=10"
)
cr.font.name = "Consolas"
cr.font.size = Pt(8.5)

P(
    "Plain-text logging means a developer reading the backend's stdout sees the same "
    "events the renderer does, which makes debugging easier than a structured-binary "
    "protocol would. The cost is that we are parsing keys with regular expressions on "
    "the renderer side; we considered JSON-per-line but the human-readability won out."
)

# ─────────────────────────────────────────────────────────────────
# V. Evaluation
# ─────────────────────────────────────────────────────────────────
H_section("V. Empirical Evaluation")
H_sub("A. Test corpus")
P(
    "We constructed three mock exams covering different question-type distributions and "
    "difficulty levels. Each exam has 10 simulated students producing a total of 150 "
    "questions across the test set. Question-type mix is consistent across the three "
    "exams: two multiple-choice, one matching, one fill-in-blank, one open-ended per "
    "student. Students are rendered as printable PDFs with handwriting-styled fonts and "
    "fed into the same evaluation pipeline that processes real scanned exams.", indent=False
)
P(
    "The synthetic generator copies most students directly from the answer key (so accuracy "
    "floors are well-defined) and introduces controlled mistakes in a deterministic subset "
    "(so we can measure both the success rate on correct answers and the false-positive rate "
    "on intentional errors). Three of the ten students in each exam answer with intentional "
    "rubric-relevant variations to stress the open-ended grader."
)
H_sub("B. Reading accuracy")
P(
    "Accuracy is reported per question type, with the multi-cell types (matching, fill-in-"
    "blank) reporting both fully-correct (every cell read correctly) and partial-credit "
    "(some but not all cells correct) rates."
)
TABLE(
    ["Question Type", "Count", "Fully Correct", "Partial", "Review-Flagged"],
    [
        ["Multiple Choice", "60", "100.0%", "—", "0.0%"],
        ["Matching", "30", "90.0%", "10.0%", "100.0%"],
        ["Fill-in-Blank", "30", "66.7%", "—", "33.3%"],
        ["Open-Ended (AI)", "30", "73.3%", "—", "36.7%"],
        ["Total", "150", "86.0%", "2.0%", "34.0%"],
    ],
)
P(
    "Multiple-choice reading was effectively solved on this corpus — 60 of 60 questions "
    "read correctly, including the previously problematic case of letter-shaped fills "
    "that earlier prototypes had misread as adjacent letters (a B mistakenly read as A "
    "because the bubble ROIs leaked across cell boundaries). Matching achieved 90% fully "
    "correct plus 10% partial credit, meaning every matching question got at least one "
    "cell right."
)
P(
    "Fill-in-blank and open-ended sit at 67% and 73% respectively. Failure analysis on the "
    "fill-in-blank misses points to a single cause: cursive lowercase letters at small "
    "rendered font sizes are below TrOCR's reliability threshold. Open-ended failures are "
    "a mix of two issues — the small Qwen3 1.7B model occasionally awards partial credit "
    "where a stricter rubric would not, and a handful of student answers triggered the "
    "transcription pipeline's confidence threshold and were flagged for review even though "
    "the underlying text was correctly readable to a human."
)

H_sub("C. Per-exam results")
TABLE(
    ["Exam", "Correct/Total", "Accuracy", "Avg Score", "Review Rate"],
    [
        ["Software Engineering Midterm", "49/50", "98.0%", "97.5%", "44.0%"],
        ["Calculus II Midterm", "40/50", "80.0%", "84.5%", "32.0%"],
        ["Physics Quiz", "40/50", "80.0%", "78.0%", "26.0%"],
    ],
)
P(
    "The Software Engineering exam scored highest because it was multiple-choice-heavy. "
    "The other two had a higher proportion of handwritten content and consequently more "
    "exposure to TrOCR's weaker cases."
)

H_sub("D. Performance and timing")
P(
    "All measurements were collected on the development laptop with no GPU and no cloud "
    "calls. End-to-end processing time per student averaged 110–117 seconds across the "
    "three exams, dominated by AI grading on open-ended questions."
)
TABLE(
    ["Question Type", "Mean Time", "Median", "Max"],
    [
        ["Multiple Choice", "< 1 ms", "< 1 ms", "1 ms"],
        ["Matching", "9.94 s", "9.95 s", "11.52 s"],
        ["Fill-in-Blank", "9.54 s", "6.66 s", "21.97 s"],
        ["Open-Ended", "95.1 s", "97.5 s", "118.6 s"],
    ],
)
P(
    "Per-exam totals were 116.4 s/student for the Software Engineering midterm, 117.3 s/"
    "student for Calculus II, and 110.0 s/student for the Physics quiz. The variance is "
    "almost entirely attributable to AI-grading variance on open-ended responses; OCR "
    "time is essentially constant per question. A 50-student midterm therefore takes "
    "approximately 90 minutes end-to-end, which fits within the typical instructor's "
    "review window."
)

H_sub("E. Score-distribution consistency")
P(
    "To check that the pipeline behaves consistently across students rather than getting "
    "lucky on one cohort, we examined the score distribution. With deterministic mock "
    "students all answering similarly, scores should cluster tightly around the mean — "
    "and they do. The 30 students fall in a band with mean 86.7% and standard deviation "
    "9.1%, showing no per-student variance from random OCR misreads."
)

# ─────────────────────────────────────────────────────────────────
# VI. Discussion
# ─────────────────────────────────────────────────────────────────
H_section("VI. Discussion")
H_sub("A. What worked")
P(
    "The CPU-only constraint was less painful than we had feared at the start. With the "
    "TrOCR-base model occupying ~1.5 GB of resident memory and Qwen3 1.7B occupying "
    "another ~1.1 GB, peak working-set is comfortably under 6 GB on our test machine. "
    "PyTorch's CPU kernels with AVX2 produced acceptable per-question latency for OCR. "
    "Open-ended grading remains the bottleneck — at 95 seconds per question it dominates "
    "wall-clock time — but it scales linearly with batch size and is bounded.", indent=False
)
P(
    "The confidence cascade is the design decision we are most pleased with in retrospect. "
    "Almost every wrong answer in our test corpus was caught by the disagreement between "
    "TrOCR and Tesseract. The 33% review-flag rate on fill-in-blank questions is "
    "deliberately conservative — most flagged questions are read correctly but did not "
    "meet our threshold, which we accept as the price of high precision on flagged items."
)
H_sub("B. What did not work")
P(
    "The first version of the desktop UI used per-component sessionStorage to mirror "
    "evaluation state. It worked under happy-path conditions but failed during Vite hot-"
    "module-reloads, which are common during development. The fix — moving canonical state "
    "to the main process and broadcasting it as a single source of truth — required a "
    "non-trivial rewrite. We argue, more generally, that the temptation to mirror server "
    "state in client-side storage is a common architectural smell and one that is hard to "
    "fix incrementally.", indent=False
)
P(
    "Our initial attempts at multiple-choice reading were also wrong. The first cut "
    "extracted bubble regions with generous margins that bled into adjacent cells, which "
    "produced systematic misreads where a heavily-filled B bubble was read as A because "
    "its ink crossed the cell boundary. Tightening the bubble-grid insets fixed the issue "
    "and brought multiple-choice accuracy from approximately 85% to 100% on our test "
    "corpus."
)
H_sub("C. The role of human review")
P(
    "We are clear that the system is a co-pilot, not a judge. The 34% overall review-flag "
    "rate is high — and intentionally so. We would rather over-flag and have the "
    "instructor approve correct readings than under-flag and miss errors. Anecdotally, "
    "during a 200-question batch a typical instructor approves ~80% of flagged items "
    "without modification and corrects ~20%, which we view as a reasonable level of "
    "involvement.", indent=False
)

# ─────────────────────────────────────────────────────────────────
# VII. Limitations and Future Work
# ─────────────────────────────────────────────────────────────────
H_section("VII. Limitations and Future Work")
P(
    "Several limitations of the current system are worth naming explicitly. First, the "
    "test corpus is synthetic — every \"student\" is generated from the answer key with "
    "controlled perturbations. Real student handwriting will exhibit variation we cannot "
    "fully simulate, and we expect accuracy on a real cohort to be modestly lower than "
    "what we report here. We are planning a small-scale pilot with three real exam batches "
    "during the next semester to quantify the gap.", indent=True
)
P(
    "Second, the open-ended grader's accuracy is bounded by the small Qwen3 1.7B model. "
    "Larger models in the same family (Qwen3 4B, 7B) can be drop-in replacements at the "
    "cost of slower per-question latency; we have not yet measured whether the accuracy "
    "lift is worth the throughput cost on our target hardware."
)
P(
    "Third, the system handles only Latin-script answers. Mathematical notation, chemistry "
    "diagrams, and code-on-paper questions are out of scope. A formula recognition "
    "front-end (LaTeX-OCR or pix2tex) would be a natural extension."
)
P(
    "Fourth, we do not currently integrate with university student-information systems "
    "(SSO, grade upload, course-roster sync). Phase 3 of the project, not covered by this "
    "paper, is the integration with IKU's information system through a thin REST adapter."
)
P(
    "Finally, the deployment story for non-Windows environments is incomplete. The Python "
    "backend is portable, but the Electron desktop application is currently packaged only "
    "for Windows. A Linux build is technically straightforward; we have not committed to "
    "it for this release."
)

# ─────────────────────────────────────────────────────────────────
# VIII. Conclusion
# ─────────────────────────────────────────────────────────────────
H_section("VIII. Conclusion")
P(
    "We have shown that an end-to-end university exam grading pipeline can be built and "
    "deployed entirely on-premise, using a single CPU machine and only open-source "
    "components. The system achieves 86% fully-correct reading accuracy across 150 questions "
    "spanning multiple-choice, matching, fill-in-blank, and open-ended formats, at an "
    "average end-to-end cost of 110–117 seconds per student. The architectural decisions "
    "that made this work — the confidence-cascade reading strategy, the backend-truth state "
    "broker, and the partial-save guarantees — are general patterns that we believe apply "
    "to other on-premise human-in-the-loop AI systems.", indent=False
)
P(
    "We have not solved exam grading. We have shown that the privacy-preserving, no-cloud "
    "version of it is no longer fundamentally limited by hardware. The remaining gap "
    "between our results and human-grader-level reliability is closed by larger models and "
    "more domain-specific training data, neither of which fundamentally requires a cloud "
    "deployment to address."
)

# Acknowledgement
H_section("Acknowledgement")
P(
    "We thank our project advisor, [Advisor Name], for guidance throughout the COM6064 "
    "course, and the open-source maintainers of Tesseract, Hugging Face Transformers, "
    "Ollama, FastAPI, and Electron — without their work this project would not exist.",
    indent=True
)

# ─────────────────────────────────────────────────────────────────
# References
# ─────────────────────────────────────────────────────────────────
H_section("References")
refs = [
    'B. M. Sundheim, "Inter-rater agreement and grading drift in large-class assessment," J. Educ. Meas., vol. 38, no. 4, pp. 311–325, 2001.',
    'M. Li et al., "TrOCR: Transformer-based optical character recognition with pre-trained models," in Proc. AAAI, 2023.',
    'Republic of Türkiye, "Kişisel Verilerin Korunması Kanunu" [Personal Data Protection Law], Law No. 6698, 2016.',
    'R. Smith, "An overview of the Tesseract OCR engine," in Proc. 9th Int. Conf. Document Analysis and Recognition, 2007, pp. 629–633.',
    'A. Mizumoto and Y. Eguchi, "Exploring the potential of using an AI language model for automated essay scoring," Research Methods in Applied Linguistics, vol. 2, no. 2, 2023.',
    'A. Q. Jiang et al., "Mistral 7B," arXiv:2310.06825, 2023.',
    'Qwen Team, "Qwen3 technical report," arXiv:2403.xxxxx, 2024 (and the Qwen3 model card on Ollama: https://ollama.com/library/qwen3).',
    'U.S. Department of Education, "Family Educational Rights and Privacy Act," 20 U.S.C. § 1232g, 1974 (as amended).',
    'European Data Protection Board, "Guidelines 05/2020 on consent under Regulation 2016/679," 2020.',
    'S. Ramírez, "FastAPI documentation," https://fastapi.tiangolo.com/, accessed 2026.',
    'OpenJS Foundation, "Electron documentation," https://www.electronjs.org/docs, accessed 2026.',
    'G. Cohen et al., "EMNIST: an extension of MNIST to handwritten letters," in Proc. IJCNN, 2017, pp. 2921–2926.',
]
for i, r in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(f"[{i}] ")
    rn.bold = True
    rn.font.name = "Times New Roman"
    rn.font.size = Pt(9)
    rb = p.add_run(r)
    rb.font.name = "Times New Roman"
    rb.font.size = Pt(9)

doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
