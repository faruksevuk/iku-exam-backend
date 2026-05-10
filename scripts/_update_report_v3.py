"""Append a Section 8 "Reliability & UX Hardening" to Report_Faruk.docx
covering the work done after the edge-case test:

  - Backend resolver hardening (the stray-folder bug we hit)
  - Friendly backend folder picker on first failure
  - Settings → Backend Directory inline folder picker
  - Local xlsx export (exceljs) — kills the "fetch failed" path
  - Per-student progress UI restored, derived from BatchState
  - Defensive IPC error surfacing so silent buttons stop happening

The report keeps its existing typography conventions (IKU red H1s with
bottom rules, navy H2s) — we just append at the end. Old sections are
untouched.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn
from docx.shared import Cm, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SOURCE = Path(r"C:/Users/faruk/Downloads/Report_Faruk.docx")
OUT = SOURCE  # overwrite in place

ACCENT = RGBColor(0xED, 0x1B, 0x24)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_color)
    tc_pr.append(shd)


doc = Document(str(SOURCE))


def H1(text: str):
    p = doc.add_heading("", level=1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def H2(text: str):
    p = doc.add_heading("", level=2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def P(text, *, italic=False, justify=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def TABLE(headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths_cm:
        t.autofit = False
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "ED1B24")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        zebra = ri % 2 == 0
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            if zebra:
                set_cell_shading(cell, "FAFBFC")


# ── Section 8 ─────────────────────────────────────────────────
doc.add_page_break()
H1("8. Reliability & UX Hardening")

P(
    "After the edge-case test landed I spent a session on the seams "
    "between the desktop app and the backend — the parts a teacher "
    "actually touches every day. Several of these came out of issues "
    "we hit ourselves during the edge-case work: a backend that "
    "wouldn't start, a download button that threw \"fetch failed\", a "
    "progress panel that didn't show enough. The fixes are small but "
    "they meaningfully change the experience of running this on a "
    "lab machine."
)

# 8.1
H2("8.1 Backend folder resolver — hardening + folder picker")

P(
    "The Electron app spawns the Python backend on launch by probing "
    "common parent folders for one that contains app.py. The first "
    "version of this only checked for directory existence, which "
    "blew up the first time we ran a bulk-export job: the export "
    "code unconditionally created a sibling \"iku-exam-backend/"
    "samples/blanks/\" directory (using mkdir-recursive), and from "
    "then on the resolver matched that empty parent first and "
    "uvicorn died with \"Could not import module 'app'\" on every "
    "restart attempt."
)

P(
    "Two fixes shipped:"
)

P(
    "First, the resolver now requires each candidate folder to "
    "contain app.py — directory existence alone isn't enough. So a "
    "stray sibling can't shadow the real one. The bulk-export code "
    "was also fixed to use the already-resolved backend dir instead "
    "of hardcoding the path, so it can't recreate the same orphan."
)

P(
    "Second, when the resolver still can't find a backend (typical "
    "case for a teammate who cloned the two repos to arbitrary "
    "parent folders), the app now shows a friendly dialog: "
    "\"Couldn't locate the FastAPI backend automatically. Point the "
    "app at the folder that contains app.py — the choice is saved.\" "
    "If the teammate picks a wrong folder, an inline error explains "
    "why before re-opening the picker. The chosen path is persisted "
    "to settings.json, so subsequent launches are silent."
)

P(
    "The same folder picker is available from two settings entry "
    "points after the fact: the AI-status pill's Backend Settings "
    "modal (a Browse… button next to the path field), and the App "
    "Settings → Backend Directory section (an inline picker that "
    "replaces the previous \"Open backend settings…\" link). "
    "Nobody has to type a path."
)

# 8.2
H2("8.2 Excel export — moved out of the backend round-trip")

P(
    "The Download Excel button used to look for a pre-generated "
    "xlsx file on disk (written by /evaluate at the end of a run), "
    "and fell back to GET /results/{id}/excel on the backend if the "
    "file was missing. Both paths were brittle: the on-disk file "
    "became stale the moment a teacher overrode a single grade, and "
    "the HTTP fallback threw a generic \"fetch failed\" any time "
    "the dynamic backend port rotated (which it does on restart)."
)

P(
    "The xlsx generation is a formatting concern. Results.json is "
    "the canonical state and already includes every teacher "
    "override. So the cleanest fix was to port the Python "
    "export.export_results() logic to TypeScript on the renderer "
    "side, using the exceljs library. Schema parity with the Python "
    "version is preserved (same Scores + Summary sheets, same "
    "colour palette for correct / partial / wrong / review)."
)

P(
    "Outcomes: Excel download now mirrors CSV — zero backend "
    "dependency, always reflects the latest overrides, works on "
    "older evaluations whose results.json predates the current run. "
    "Tested by generating a workbook for an 11-student record with "
    "no backend running and confirming Scores sheet rendered "
    "correctly with proper conditional fills."
)

# 8.3
H2("8.3 Per-student progress UI — restored and rebuilt")

P(
    "Earlier in the project we had a per-student row list during a "
    "run. It was removed because the implementation kept local "
    "session-storage mirrors of progress state, which drifted out "
    "of sync after Vite hot-reload and showed the dreaded \"AI "
    "grading Q5\" stuck panel forever. The architectural fix for "
    "that was the BatchState broker (the main process is the single "
    "source of truth), but the visual surface was never rebuilt."
)

P(
    "This sprint I rebuilt it correctly. The broker now carries a "
    "completedStudents[] array that grows each time the pipeline "
    "emits stage=done — one entry per finished student with "
    "{sIdx, sNum, hadAi, ocrMs, aiMs, finishedAt}. The renderer "
    "renders three buckets purely from the broker state:"
)

TABLE(
    ["Bucket", "Row state", "Source"],
    [
        ["Done", "Green ✓, OCR'd SN, optional [AI] badge, \"done\"",
         "batchState.completedStudents (one row per finished student)"],
        ["Current", "Pulsing red ●, live SN, live Q + question-type label",
         "batchState.sIdx + sNum + qNum + qType"],
        ["Pending", "Dashed ○, no SN, \"pending\"",
         "Slots i ≥ completedStudents.length up to sTotal"],
    ],
    col_widths_cm=[2.4, 6.0, 8.6],
)

P(
    "No local state, no sessionStorage, no remount drift by "
    "construction. For huge cohorts (more than 10 students) only a "
    "window of rows around the active one is shown, with \"+N "
    "earlier\" and \"+N more pending\" overflow chips collapsing "
    "the rest so the card doesn't stretch infinitely."
)

# 8.4
H2("8.4 Defensive IPC — surfacing silent failures")

P(
    "An anti-pattern I removed while doing this work: several IPC "
    "callers in the renderer had the shape `if (!api?.method) "
    "return`. The intent was \"don't crash if the bridge is "
    "missing,\" but the effect was that a button click silently did "
    "nothing — no toast, no console error, no log line. We saw this "
    "directly during the folder-picker rollout: I'd added the IPC "
    "but the user was running an older main process that hadn't "
    "registered the handler yet, the invoke rejected, and the click "
    "produced zero visible feedback."
)

P(
    "Fix: every IPC caller in Settings now distinguishes three "
    "failure modes and surfaces each as inline UI feedback — (1) "
    "missing electronAPI entirely (\"open via the app, not a "
    "browser\"); (2) defined preload but missing channel handler "
    "(\"This feature was added in a newer build — please close and "
    "reopen the app\"); (3) handler registered but the invoke "
    "rejects (\"Folder picker failed: <message>. Try restarting "
    "the app\"). Buttons that look like they did nothing now tell "
    "you precisely why."
)

# 8.5
H2("8.5 Summary — what changed under the hood")

TABLE(
    ["Area", "Before", "After"],
    [
        ["Backend folder lookup",
         "Plain existsSync; one stray sibling could shadow the real backend; manual env var or hand-typed path to recover.",
         "Validates app.py exists; auto-dialog on first failure with folder picker; persists to settings.json; picker also available from two settings entry points."],
        ["Excel download",
         "Read pre-generated xlsx from disk; HTTP fallback to backend; both could fail (stale, missing, port rotated).",
         "Generated locally via exceljs from the live results.json; zero backend dependency; always reflects overrides."],
        ["Live evaluation progress",
         "Single \"Reading Q3 …\" line + counters.",
         "Per-student row list with ✓ done / ● current / ○ pending states, derived purely from BatchState.completedStudents + sIdx/sNum/qNum/qType."],
        ["IPC failure modes",
         "Silent — null-guard returns; button clicks could do nothing.",
         "Three explicit failure paths, each with actionable inline error text."],
        ["Bulk export hardcoded path",
         "Always wrote to <app>/../iku-exam-backend/samples/blanks/ — could create a phantom directory that broke the resolver.",
         "Uses the resolved backend dir; generic fallback only when no resolved dir is available."],
    ],
    col_widths_cm=[3.5, 6.5, 7.0],
)

P(
    "None of these changes touched the OCR / OMR / LLM grading paths "
    "— those modules are out of scope for this round and remain on "
    "the accuracy numbers from Sections 2 and 7. What changed is "
    "the interface layer: how a teacher reaches those paths, "
    "watches them run, and exports the result. The point of this "
    "sprint was to make all of that resilient enough that a "
    "teammate could clone the two repos, click the icon, and not "
    "get stuck on a red status pill."
)

doc.save(str(OUT))
print(f"Updated report: {OUT}  ({OUT.stat().st_size / 1024:.1f} KB)")
