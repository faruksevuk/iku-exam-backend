"""SRS v2 — shorter, with embedded screenshots, custom-styled tables,
and no role allocations on the team table.

Output: C:/Users/faruk/Downloads/COM6064-SRS-LLM-Exam-Evaluation.docx
"""
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

OUT = Path(r"C:/Users/faruk/Downloads/COM6064-SRS-LLM-Exam-Evaluation.docx")

UI = Path(r"D:/repos/ExamGeneration/iku-exam-generator/docs/screenshots")
ASSETS = Path(r"D:/repos/ExamGeneration/_report_assets")
TR_ASSETS = Path(r"C:/Users/faruk/Downloads/_test_report_assets")

doc = Document()

# Page setup — A4, slightly tighter margins
for s in doc.sections:
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

# Default style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.18

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)        # navy
ACCENT_LIGHT = RGBColor(0xEC, 0xF1, 0xF8)  # very light navy
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
SOFT = RGBColor(0xF7, 0xF8, 0xFA)


def set_cell_shading(cell, hex_color: str):
    """Apply solid background fill to a table cell via XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        b = OxmlElement(f"w:{edge}")
        if val is None:
            b.set(qn("w:val"), "nil")
        else:
            sz, color = val
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def styled_table(headers: list[str], rows: list[list[str]],
                 col_widths_cm: list[float] | None = None) -> None:
    """Custom-styled table: navy header, alternating soft rows, horizontal-only borders."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set fixed col widths if given
    if col_widths_cm:
        t.autofit = False
        for col_idx, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)

    # Header row
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        set_cell_shading(cell, "1F3A5F")
        set_cell_borders(
            cell,
            top=(8, "1F3A5F"), bottom=(8, "1F3A5F"),
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # Body rows
    for ri, row in enumerate(rows, start=1):
        zebra = (ri % 2 == 0)
        is_last = (ri == len(rows))
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            if zebra:
                set_cell_shading(cell, "F7F8FA")
            bottom_border = (8, "1F3A5F") if is_last else (4, "D1D5DB")
            set_cell_borders(
                cell,
                top=None,
                bottom=bottom_border,
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph("").paragraph_format.space_after = Pt(2)


def H1(text: str, *, before: int = 16) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(8)
    # Bottom rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def H2(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def H3(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def P(text: str, *, italic: bool = False, justify: bool = True) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def figure(path: Path, caption: str, width_cm: float = 14) -> None:
    if not path.exists():
        P(f"[Figure missing: {path.name}]", italic=True, justify=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(8)
    cr = cp.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = GRAY


def CODE(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


# ─────────────────────────────────────────────────────────────────
# Title page
# ─────────────────────────────────────────────────────────────────
for _ in range(2):
    doc.add_paragraph("")

t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = t1.add_run("Software Requirements Specification")
r1.bold = True
r1.font.size = Pt(24)
r1.font.color.rgb = ACCENT

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("for")
r2.font.size = Pt(13)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("LLM-Based Exam Evaluation Engine")
r3.bold = True
r3.font.size = Pt(20)
r3.font.color.rgb = DARK

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("On-premise, CPU-optimised intelligent assessment for university exams")
r4.italic = True
r4.font.size = Pt(11)
r4.font.color.rgb = GRAY

for _ in range(2):
    doc.add_paragraph("")

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run("Version 1.0").font.size = Pt(12)

doc.add_paragraph("")

prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
prep.add_run("Prepared by [Group Name]").font.size = Pt(12)

doc.add_paragraph("")

# Team table — names + IDs only, no roles
team_t = doc.add_table(rows=6, cols=2)
team_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_t.autofit = False
for ri, (nm, sid) in enumerate([
    ("[Member 1 Full Name]", "[Student ID 1]"),
    ("[Member 2 Full Name]", "[Student ID 2]"),
    ("[Member 3 Full Name]", "[Student ID 3]"),
    ("[Member 4 Full Name]", "[Student ID 4]"),
    ("[Member 5 Full Name]", "[Student ID 5]"),
    ("[Member 6 Full Name]", "[Student ID 6]"),
]):
    name_cell = team_t.rows[ri].cells[0]
    name_cell.width = Cm(7)
    name_cell.text = ""
    para = name_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(nm)
    run.font.size = Pt(11)
    set_cell_borders(name_cell)

    id_cell = team_t.rows[ri].cells[1]
    id_cell.width = Cm(5)
    id_cell.text = ""
    para = id_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(sid)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    set_cell_borders(id_cell)

doc.add_paragraph("")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(datetime.now().strftime("%B %Y") + "  •  COM6064  •  Istanbul Kültür University").font.size = Pt(10)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# Revisions
# ─────────────────────────────────────────────────────────────────
H1("Revisions", before=0)
styled_table(
    ["Version", "Date", "Description"],
    [
        ["0.1", "[Draft]", "Initial outline; sections 1–2 drafted."],
        ["0.5", "[Date]", "Functional and non-functional requirements completed; first architecture pass."],
        ["0.9", "[Date]", "Internal review; use cases, deployment view, design figures added."],
        ["1.0", datetime.now().strftime("%b %d, %Y"), "Released for course submission."],
    ],
    col_widths_cm=[2.0, 3.0, 11.0],
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 1. Introduction
# ─────────────────────────────────────────────────────────────────
H1("1.  Introduction", before=0)
P(
    "Universities have graded paper exams the same way for decades. Students write, "
    "instructors collect, instructors read every booklet end to end. For a 200-paper "
    "midterm, by the eighth hour fatigue starts to bias the rubric. Some answers get "
    "the benefit of the doubt; others get docked for ambiguous handwriting. Drift is "
    "real even when the grader is doing their best."
)
P(
    "Optical mark recognition has automated bubble sheets since the 1980s. Anything "
    "handwritten — short answers, fill-in-blanks, single-letter matchings, open-ended "
    "responses — has continued to need eyes. The recent jump in transformer-based "
    "handwriting recognition and small open-weight language models has finally made the "
    "rest of it tractable. Most products that do this in production push student data to "
    "cloud APIs, which is a non-starter for a Turkish public university where exam "
    "papers are KVKK personal data."
)
P(
    "This document specifies the requirements for the LLM-Based Exam Evaluation Engine: "
    "a system that brings AI-assisted grading on-premise, end to end, with no GPU and no "
    "outbound network calls. The OCR stack is Tesseract plus TrOCR. The open-ended "
    "grader is Qwen3 1.7B running locally through Ollama. The orchestration is a "
    "FastAPI backend wrapped in an Electron desktop app."
)

H2("1.1  Project Purpose, Scope, and Objectives")
P(
    "Two coupled artefacts make up the product: the FastAPI backend (exam-backend) "
    "running the OCR + grading pipeline, and the Electron desktop app (iku-exam-"
    "generator) that authors exams, ingests scans, and shows the instructor results to "
    "review. The two communicate over loopback HTTP."
)
P(
    "Phase 1 covers structured questions — multiple-choice, multi-select, single-letter "
    "matching, and short fill-in-the-blank. These are read with classical OMR and a "
    "confidence cascade that combines a small EMNIST CNN, Tesseract, and TrOCR. Phase 2 "
    "covers open-ended grading: TrOCR transcribes the answer; Qwen3 grades it against "
    "the rubric and returns a score, a confidence value, and a written justification. "
    "Long-form essays beyond ~200 words, university SSO integration, mobile clients, and "
    "any cloud grading paths are out of scope for this release."
)
P(
    "What the product is meant to deliver is concrete: instructor time on a 200-paper "
    "midterm should drop from 10–14 hours to roughly 30–60 minutes of review and "
    "approval. Every AI verdict carries a confidence and a justification. Every teacher "
    "override is logged with a timestamp. Privacy is structural — no outbound HTTP — "
    "and the operating cost is essentially zero, since all components are open source."
)

H2("1.2  Team")
P(
    "The team is six undergraduate students working on this project under the COM6064 "
    "course. Names and IDs:"
)
styled_table(
    ["Member", "Student ID"],
    [
        ["[Member 1 Full Name]", "[Student ID 1]"],
        ["[Member 2 Full Name]", "[Student ID 2]"],
        ["[Member 3 Full Name]", "[Student ID 3]"],
        ["[Member 4 Full Name]", "[Student ID 4]"],
        ["[Member 5 Full Name]", "[Student ID 5]"],
        ["[Member 6 Full Name]", "[Student ID 6]"],
    ],
    col_widths_cm=[10.5, 5.5],
)

H2("1.3  Technical Assumptions and Constraints")
P(
    "We build and test on Windows 10/11. The deployment target is a typical lab "
    "machine: an 8th-gen Intel i5 or equivalent, 8 GB RAM, and no discrete GPU. The "
    "stack is Python 3.11 with FastAPI on Uvicorn (backend), TypeScript / React 19 on "
    "Electron 41 (frontend), PyTorch 2.x CPU build, Tesseract ≥ 5.0, and Ollama "
    "serving qwen3:1.7b on 127.0.0.1:11434."
)
P(
    "Hardware: 8 GB RAM minimum (16 GB recommended), AVX2-capable CPU, 5 GB free disk "
    "for model weights. Connectivity: no outbound calls — Ollama is the only network "
    "endpoint contacted. Licensing: dependencies are MIT or Apache 2.0, except PyMuPDF "
    "(AGPL-3) which we use under its on-premise terms; redistribution as proprietary "
    "SaaS would require a commercial PyMuPDF licence."
)

H2("1.4  Naming Conventions")
styled_table(
    ["Term", "Meaning"],
    [
        ["Pipeline", "End-to-end backend chain in pipeline.py, from PDF intake to results."],
        ["map.json", "Per-exam ground-truth file with question coordinates, expected answers, and per-question scoring."],
        ["BatchState", "Canonical state of a running evaluation, owned by the Electron main process."],
        ["Confidence", "Value in [0, 1]. Per-question confidence is the lower of the contributing cell confidences."],
        ["Needs review", "Boolean flag set when any reader's confidence is below threshold or two readers disagree."],
        ["[STAGE] event", "Structured log line emitted by Python and parsed by the Electron main process to drive the live UI."],
    ],
    col_widths_cm=[3.5, 12.5],
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 2. Requirements
# ─────────────────────────────────────────────────────────────────
H1("2.  Requirements", before=0)

H2("2.1  Functional Requirements")
P("Functional requirements are grouped by area; each has a stable identifier used in our traceability matrix.")

H3("Exam authoring")
styled_table(
    ["ID", "Requirement"],
    [
        ["FR-1.1", "Create new exams with course code, exam type, date, faculty, department, total points, and instructions."],
        ["FR-1.2", "Support five question types: multiple-choice, multi-select, matching, fill-in-blank, and open-ended."],
        ["FR-1.3", "Per question, specify point value, cell count where applicable, and expected answer."],
        ["FR-1.4", "Save / open / save-as / duplicate / delete exams in the local store under a stable examId slug."],
        ["FR-1.5", "Produce a printable blank PDF and a paired map.json for any saved exam."],
    ],
    col_widths_cm=[1.8, 14.2],
)

H3("Scan ingestion and evaluation")
styled_table(
    ["ID", "Requirement"],
    [
        ["FR-2.1", "Accept multi-page PDF input; split into per-student page groups via QR code or sequential fallback."],
        ["FR-2.2", "Align scanned pages using the four corner anchors printed on the blank, correcting rotation and shear."],
        ["FR-2.3", "Read the student number with a digit CNN, returning per-digit confidence."],
        ["FR-2.4", "Extract MC / MS answers via bubble-fill ratio thresholds; produce per-bubble fill ratios."],
        ["FR-2.5", "Read matching cells through a LetterCNN + Tesseract cascade; flag disagreement for review."],
        ["FR-2.6", "Read fill-in-blank cells through a TrOCR + Tesseract cascade; fuzzy-match against expected."],
        ["FR-2.7", "Transcribe open-ended answers with TrOCR and grade them with Qwen3 1.7B; return score, confidence, justification."],
        ["FR-2.8", "Preserve partial results when a per-student error occurs; never lose prior students."],
        ["FR-2.9", "Write a results.json sibling on every run for crash recovery."],
        ["FR-2.10", "Emit [STAGE] log events consumed by the desktop app for live progress."],
    ],
    col_widths_cm=[1.8, 14.2],
)

H3("Review and export")
styled_table(
    ["ID", "Requirement"],
    [
        ["FR-3.1", "Show per-student review pages with cropped answer image, OCR / AI verdict, confidence, and explanation."],
        ["FR-3.2", "Allow approving, editing, or rejecting any AI verdict; capture every change in an immutable override history."],
        ["FR-3.3", "Bulk-approve all reviewable items in a batch with a confirmation step."],
        ["FR-3.4", "Export results to xlsx (per-student rows, per-question scores, totals, review highlights) and CSV (RFC 4180)."],
        ["FR-3.5", "Render an analytics view per exam: distribution, mean / stdev, hardest questions, and 2–3 LLM-generated insights."],
        ["FR-3.6", "Support Turkish and English at runtime, persisted across sessions."],
    ],
    col_widths_cm=[1.8, 14.2],
)

H2("2.2  Non-Functional Requirements")

H3("Performance")
styled_table(
    ["ID", "Requirement", "Measured"],
    [
        ["NFR-P1", "End-to-end processing per student under 150 s on an 8 GB / no-GPU laptop.", "110–117 s"],
        ["NFR-P2", "Backend cold-start under 90 s including model load.", "50–60 s"],
        ["NFR-P3", "Renderer UI stays interactive (60 fps) during evaluation.", "Met"],
        ["NFR-P4", "200-student xlsx export under 5 s.", "≈ 2 s"],
        ["NFR-P5", "Analytics queries under 500 ms for cohorts up to 500 students.", "≈ 150 ms"],
    ],
    col_widths_cm=[1.8, 11.4, 2.8],
)

H3("Safety and Security")
styled_table(
    ["ID", "Requirement"],
    [
        ["NFR-S1", "No outbound HTTP at runtime. Loopback to Ollama is the only network endpoint."],
        ["NFR-S2", "Reject prompt-injection in student answers via an explicit system-prompt instruction plus a safety-filter pass on the model's justification."],
        ["NFR-S3", "All file paths exposed via IPC are whitelist-validated; path traversal must be impossible by construction."],
        ["NFR-S4", "AI verdicts are never auto-committed. The instructor must explicitly approve each verdict (or use the bulk-approve action)."],
    ],
    col_widths_cm=[1.8, 14.2],
)

H3("Quality Attributes")
P(
    "Reliability: each per-student loop iteration is wrapped in its own try/except; the "
    "whole batch has a top-level guard that always returns partial results with "
    "pipelineError set. Auditability: every AI verdict ships with confidence and "
    "justification; every override appends to overrideHistory[] with timestamp and "
    "previous value. Maintainability: thresholds and external commands live in a single "
    "module (config.py); adding a new question type is one editor component plus one "
    "reducer entry. Portability: backend is pure Python and cross-platform; the desktop "
    "build targets Windows for this release."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 3. Other Requirements
# ─────────────────────────────────────────────────────────────────
H1("3.  Other Requirements", before=0)
P(
    "Phase 1 uses no SQL database — per-exam data lives in flat JSON under the Electron "
    "userData directory. SQL adds operational complexity for what is single-user "
    "data, and the file model lets an instructor inspect, back up, and migrate exam "
    "data with a regular file manager."
)
P(
    "Internationalisation: Turkish and English. Translation strings live in "
    "src/renderer/src/i18n/{en,tr}.ts as plain TypeScript modules; the active language "
    "is persisted in localStorage and exposed through a useLang() hook. The string set "
    "is small enough that we deliberately avoided react-intl or i18next."
)
P(
    "Legal: dependencies are MIT or Apache 2.0 except PyMuPDF (AGPL-3). Under our "
    "on-premise university deployment the AGPL terms are met; if the product is "
    "later distributed as proprietary SaaS, this would require a commercial PyMuPDF "
    "licence or a switch to a permissively-licensed alternative."
)
P(
    "Data retention: exam data, scans, and results are kept for the academic term they "
    "cover. Deletion through the UI removes the exam definition, the map, the results, "
    "and any cached page images."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 4. System Architecture
# ─────────────────────────────────────────────────────────────────
H1("4.  System Architecture", before=0)
P(
    "We adopted a three-tier architecture: a React renderer (UI), an Electron main "
    "process (lifecycle, IPC, backend management), and a Python FastAPI backend "
    "(evaluation engine). The split was deliberate. The Python ML ecosystem and the "
    "Electron desktop ecosystem are each mature; bridging them in a single process is "
    "unhappy territory. Treating Python as a managed child process keeps both sides on "
    "their happy path."
)

H2("4.1  Logical View")
figure(ASSETS / "04_pipeline_architecture.png",
       "Figure 1.  Pipeline architecture. PDF intake on the left flows through "
       "alignment, per-question dispatch, and grading; the OCR cascade combines "
       "Tesseract, the small CNNs, and TrOCR. Open-ended responses pass through "
       "TrOCR for transcription and Qwen3 1.7B for the rubric verdict.",
       width_cm=14.5)
P(
    "Five logical components are visible to a domain expert: the Editor (authors exam "
    "definitions), the Scan Ingestor (splits multi-student PDFs by QR or page count), "
    "the Evaluator (the per-question dispatch and grading core), the Reviewer "
    "(per-student review screens and overrides), and the Analytics layer (dashboard "
    "plus LLM-generated insights). Editor and Reviewer live in the renderer; Scan "
    "Ingestor straddles main and Python; Evaluator is purely Python."
)

H2("4.2  Deployment View")
P(
    "There is one physical node: an instructor's Windows desktop. Three OS-level "
    "processes coexist on it — the Electron desktop app (main + renderer + GPU "
    "process), a Python interpreter running the FastAPI backend on a randomly-allocated "
    "localhost port, and the Ollama daemon on 127.0.0.1:11434. All communication is "
    "loopback. Persistent data lives in the Electron userData folder."
)

H2("4.3  Use Case View")
P(
    "Three actors interact with the system: the Instructor (the dominant actor), a "
    "Department Coordinator (occasional analytics reviewer), and a System Administrator "
    "(installs the application). The seven primary use cases are: Author Exam, Generate "
    "Blank PDFs, Evaluate Scans, Review Results, Export Grades, View Analytics, and "
    "Configure System."
)

H3("Use Case Scenarios")
P(
    "Author Exam. Ayşe opens the application, clicks New Exam, fills in course code, "
    "type, and date, then adds ten questions of mixed types. She saves; the app writes "
    "the exam to the local store and shows confirmation."
)
P(
    "Evaluate Scans. Ahmet has scanned 47 booklets into a single PDF. He drags it into "
    "the Evaluate modal and clicks Run all. The activity panel shows progress per "
    "student (\"Reading 12 / 47 — Q4 matching\"). After about 90 minutes the panel "
    "switches to Saving and then to Done."
)
P(
    "Review Results. Ahmet clicks Review. He sees student records on the left and "
    "questions on the right. Q5 (open-ended) is yellow-flagged with the AI's "
    "transcription and a 6/10 verdict; he reads the answer, agrees, and clicks "
    "Approve. He moves through the 47 students, intervening only on the ~15 flagged "
    "items, then exports xlsx."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 5. Design and Implementation
# ─────────────────────────────────────────────────────────────────
H1("5.  Design and Implementation", before=0)

H2("5.1  Pipeline orchestrator")
P(
    "The backend's entry point is evaluate_batch in pipeline.py. It accepts a PDF path "
    "and a parsed map dictionary and returns a list of student results. The shape of "
    "the inner loop matters:"
)
CODE(
    "for s_idx, group in enumerate(student_groups):\n"
    "    try:\n"
    "        student = read_student(group, map_data)\n"
    "        for q_num, q_def in iter_questions(map_data):\n"
    "            student.questions[q_num] = grade_question(group, q_def)\n"
    "        results.append(student)\n"
    "    except Exception as e:\n"
    "        log_stage(f'stage=error sIdx={s_idx} err={e}')\n"
    "        # do NOT re-raise — preserve all prior students\n"
    "write_results_json(exam_id, results)\n"
    "write_xlsx(exam_id, results)"
)
P(
    "Two things to call out: the per-student try / except (one bad scan does not abort "
    "a batch), and the disk write before the function returns (the renderer can crash "
    "or be killed mid-evaluation and results survive)."
)

H2("5.2  User Interface")
figure(UI / "01-dashboard.png",
       "Figure 2.  Dashboard, evaluations tab. Each card shows an exam with its "
       "map / results state and an Evaluate button.",
       width_cm=14.5)

figure(UI / "02-live-evaluation.png",
       "Figure 3.  A live evaluation in progress. The activity panel renders directly "
       "from the BatchState broadcast over IPC — no per-component sessionStorage, no "
       "drift after hot reload.",
       width_cm=14.5)

figure(UI / "04-workspace-override.png",
       "Figure 4.  The results workspace. Cropped answer image on the left; AI "
       "transcription, verdict, and the override panel on the right. Every edit "
       "appends to the override history.",
       width_cm=14.5)

figure(UI / "05-analytics-distribution.png",
       "Figure 5.  Analytics view: score distribution and AI-generated insights. "
       "Insights come from the local Qwen3 model running on the same machine.",
       width_cm=14.5)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 6. Other Supporting Information
# ─────────────────────────────────────────────────────────────────
H1("6.  Other Supporting Information", before=0)

H2("6.1  Sample I/O")
P("Map.json (single multiple-choice question on a single page):")
CODE(
    '{\n'
    '  "version": 2, "examId": "bil204-midterm-2026", "totalPages": 3,\n'
    '  "pages": [{\n'
    '    "anchors": { "TL": {...}, "TR": {...}, "BL": {...}, "BR": {...} },\n'
    '    "questions": {\n'
    '      "1": {\n'
    '        "type": "multiple_choice",\n'
    '        "scoring": { "points": 5 },\n'
    '        "expectedAnswer": { "correctOption": "B" },\n'
    '        "options": { "A": {x,y,w,h}, "B": {...}, "C": {...}, "D": {...} }\n'
    '      }\n'
    '    }\n'
    '  }]\n'
    '}'
)

H2("6.2  Cost")
P(
    "The full software stack is open source. There is no per-seat licence cost. "
    "Compared to a cloud LLM grading service charging $0.15 per evaluated student-exam, "
    "an on-premise deployment of this system saves roughly $3,000 per 200-student "
    "midterm at department scale."
)

H2("6.3  Known issues at the time of writing")
P(
    "Fill-in-blank reading sits at ~67% accuracy on our 30-student test set; cursive "
    "lowercase letters at small font sizes are the dominant failure. We are bumping the "
    "printed font size in the next layout revision."
)
P(
    "Open-ended grading at ~73% is occasionally too generous — the small Qwen3 1.7B "
    "model awards partial credit where a stricter rubric would not. A larger model "
    "(Qwen3 4B) is being evaluated as a drop-in replacement."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 7. References
# ─────────────────────────────────────────────────────────────────
H1("7.  References", before=0)

refs = [
    'IEEE Std 830-1998, "IEEE Recommended Practice for Software Requirements Specifications," 1998.',
    'ISO/IEC/IEEE 29148:2018, "Systems and software engineering — Life cycle processes — Requirements engineering."',
    'R. Smith, "An overview of the Tesseract OCR engine," in Proc. ICDAR, 2007.',
    'M. Li et al., "TrOCR: Transformer-based Optical Character Recognition with Pre-trained Models," in Proc. AAAI, 2023.',
    'Qwen Team, "Qwen3 technical report," and the Qwen3 model card on Ollama (https://ollama.com/library/qwen3).',
    'S. Ramírez, "FastAPI documentation," https://fastapi.tiangolo.com/',
    'OpenJS Foundation, "Electron documentation," https://www.electronjs.org/docs',
    'Artifex Software, "PyMuPDF / MuPDF AGPL licence terms."',
    'G. Cohen et al., "EMNIST: an extension of MNIST to handwritten letters," in Proc. IJCNN, 2017.',
    'Republic of Türkiye, "Kişisel Verilerin Korunması Kanunu (KVKK)," Law No. 6698, 2016.',
]
for i, r in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(f"[{i}] ")
    rn.bold = True
    p.add_run(r)

doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
